from __future__ import annotations

from pathlib import Path

from docx import Document

from .models import ProposalSubmission
from .storage import OUTPUT_DIR, ensure_directories


BASE_DIR = Path(__file__).resolve().parents[1]
CONTRACT_TEMPLATE_PATH = BASE_DIR / "templates" / "contrato_locacao_modelo.docx"


def _fallback(value: str, fallback: str = "A preencher") -> str:
    return value.strip() if value and value.strip() else fallback


def _build_tenant_details(submission: ProposalSubmission) -> str:
    applicant = submission.applicant_data
    return (
        f"{_fallback(applicant.full_name)}, {_fallback(applicant.nationality).lower()}, "
        f"{_fallback(applicant.marital_status).lower()}, {_fallback(applicant.profession).lower()}, "
        f"portador(a) do documento {_fallback(applicant.rg_info)}, inscrito(a) no CPF/MF sob o nº "
        f"{_fallback(applicant.cpf)}, e-mail: {_fallback(applicant.email)}, WhatsApp {_fallback(applicant.phone)}, "
        f"residente e domiciliado(a) em {_fallback(applicant.current_address)}."
    )


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, value)

    if new_text == full_text:
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_placeholders(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


def generate_contract_docx(submission: ProposalSubmission) -> Path:
    ensure_directories()
    if not CONTRACT_TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Modelo de contrato nao encontrado: {CONTRACT_TEMPLATE_PATH}")

    contract = submission.contract_data
    property_data = submission.property_data
    applicant = submission.applicant_data
    output_path = OUTPUT_DIR / f"{submission.protocol}-contrato-modelo.docx"

    document = Document(CONTRACT_TEMPLATE_PATH)
    replacements = {
        "{{ENDERECO_IMOVEL}}": _fallback(property_data.address),
        "{{DADOS_LOCADORA}}": _fallback(contract.landlord_details),
        "{{DADOS_LOCATARIO}}": _build_tenant_details(submission),
        "{{PRAZO_LOCACAO}}": _fallback(contract.lease_term),
        "{{DATA_INICIO}}": _fallback(contract.start_date),
        "{{DATA_TERMINO}}": _fallback(contract.end_date),
        "{{VALOR_ALUGUEL}}": _fallback(property_data.proposed_rent),
        "{{TIPO_LOCACAO}}": _fallback(property_data.lease_type).upper(),
        "{{CLAUSULA_GARANTIA}}": _fallback(contract.guarantee_clause),
        "{{LOCAL_DATA_ASSINATURA}}": _fallback(contract.signature_location_date),
        "{{NOME_LOCADORA_ASSINATURA}}": _fallback(contract.landlord_signature_name),
        "{{NOME_REPRESENTANTE_LOCADORA}}": _fallback(contract.landlord_representative_name),
        "{{NOME_LOCATARIO_ASSINATURA}}": _fallback(applicant.full_name),
    }
    _replace_placeholders(document, replacements)
    document.save(output_path)
    return output_path
